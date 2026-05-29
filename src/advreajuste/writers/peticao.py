"""Petição via docxtpl (Jinja2 em .docx) ou geração direta via python-docx."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt


def gerar_peticao(
    saida: Path,
    contexto: dict[str, Any],
    template: Path | None = None,
) -> Path:
    """Se `template` existir (docxtpl Jinja2), renderiza. Senão gera Word simples."""
    saida.parent.mkdir(parents=True, exist_ok=True)
    if template and Path(template).exists():
        from docxtpl import DocxTemplate

        tpl = DocxTemplate(str(template))
        tpl.render(contexto)
        tpl.save(saida)
        return saida

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading("EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO", level=1)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        f"{contexto.get('autor', '[AUTOR]')}, já qualificado(a), por seu(sua) "
        "advogado(a) que esta subscreve, vem, respeitosamente, à presença de "
        "Vossa Excelência propor a presente"
    )
    doc.add_heading(
        "AÇÃO DECLARATÓRIA DE NULIDADE DE REAJUSTE C/C REPETIÇÃO DE INDÉBITO",
        level=2,
    )
    p = doc.add_paragraph()
    p.add_run(
        f"em face de {contexto.get('re', '[OPERADORA]')}, pelos fatos e fundamentos a seguir."
    )

    doc.add_heading("I – DOS FATOS", level=2)
    doc.add_paragraph(contexto.get("fatos", "[fatos do caso]"))

    doc.add_heading("II – DO DIREITO", level=2)
    doc.add_paragraph(
        "II.1 — Da natureza de FALSO COLETIVO (REsp 1.553.013/RJ, Min. Cueva): "
        "contrato com menos de 30 vidas equipara-se, para fins de reajuste, ao "
        "individual, devendo observar o teto ANS (RN 565/2022)."
    )
    doc.add_paragraph(
        "II.2 — Da variação acumulada como PRODUTÓRIO (STJ Tema 1016): "
        "o cálculo do reajuste cumulativo observa a fórmula ∏(1+rᵢ)−1, conforme "
        "pacificado no DJe 08/04/2022."
    )
    doc.add_paragraph(
        "II.3 — Da nulidade de reajuste por faixa etária a partir de 60 anos "
        "(Súmula 91/TJSP, Estatuto do Idoso art. 15 §3º)."
    )

    doc.add_heading("III – DOS PEDIDOS", level=2)
    doc.add_paragraph(
        f"Ante o exposto, requer:\n"
        f"a) tutela de urgência para imediata adequação da mensalidade ao teto ANS;\n"
        f"b) repetição do indébito de {contexto.get('total_devido', '[valor]')} "
        f"(trienal, art. 206 §3º IV CC);\n"
        f"c) correção monetária pela Tabela Prática TJSP (INPC pré-Lei 14.905/24, "
        f"SELIC líquida a partir de 30/08/2024);\n"
        f"d) inversão do ônus da prova (CDC art. 6º VIII, Súmula 608/STJ).\n"
    )

    doc.add_paragraph(
        f"\nDá-se à causa o valor de {contexto.get('valor_causa', '[valor]')}."
    )
    doc.add_paragraph("\nTermos em que, pede deferimento.")
    doc.add_paragraph(
        f"\n{contexto.get('cidade', '[Cidade]')}, {contexto.get('data', '[data]')}."
    )
    doc.add_paragraph("\n\n_______________________________________")
    doc.add_paragraph(f"{contexto.get('advogado', '[Advogado]')} — OAB/{contexto.get('oab', '[UF/num]')}")

    doc.save(saida)
    return saida
